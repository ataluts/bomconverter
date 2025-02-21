import os
import datetime
from copy import deepcopy
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import ImageFont
import re

from typedef_pe3 import PE3_typeDef                                 #класс перечня элементов

script_dirName  = os.path.dirname(__file__)                                                     #адрес папки со скриптом
script_baseName = os.path.splitext(os.path.basename(__file__))[0]                               #базовое имя модуля
template_defaultAddress = os.path.join(script_dirName, script_baseName + os.extsep + 'docx')    #адрес шаблона по-умолчанию

#Экспортирует ПЭ3 в формате docx
def export(data, address, **kwargs):
    print('INFO >> pe3-docx exporting module running with parameters:')
    print(' ' * 12 + 'output: ' +  os.path.basename(address))

    #адрес шаблона перечня элементов
    template_address = kwargs.get('template', template_defaultAddress)
    print(' ' * 12 + 'template: ' + os.path.basename(template_address))

    #параметры переноса строк в полях таблицы перечня 
    pe3Table_wordwrap_desig = kwargs.get('format_table_desig_wordwrap', True)
    pe3Table_wordwrap_label      = kwargs.get('format_table_label_wordwrap', True)
    pe3Table_wordwrap_annot = kwargs.get('format_table_annot_wordwrap', True)
    print_txt = []
    if pe3Table_wordwrap_desig: print_txt.append('Designator')
    if pe3Table_wordwrap_label: print_txt.append('Label')
    if pe3Table_wordwrap_annot: print_txt.append('Annotation')
    if len(print_txt) == 0:     print_txt.append('disabled')
    print(' ' * 12 + 'wordwrap: ' + ', '.join(print_txt))

    #параметры разделителей пре переносе строк
    pe3Table_wordwrap_delimiter_desig = kwargs.get('format_table_desig_wordwrap_delimiter', ' ')
    pe3Table_wordwrap_delimiter_label = kwargs.get('format_table_label_wordwrap_delimiter', ' ')
    pe3Table_wordwrap_delimiter_annot = kwargs.get('format_table_annot_wordwrap_delimiter', ' ')
    print(' ' * 12 + "wordwrap delimiter: Designator='" + pe3Table_wordwrap_delimiter_desig + "', Label='" + pe3Table_wordwrap_delimiter_label + "', Annotation='" + pe3Table_wordwrap_delimiter_annot + "'")

    #свойства документа, если не заданы то остаются как в шаблоне
    document_author         = kwargs.get('document_author', None)
    document_category       = kwargs.get('document_category', None)
    document_comments       = kwargs.get('document_comments', None)
    document_contentStatus  = kwargs.get('document_contentStatus', None)
    document_created        = kwargs.get('document_created', None)
    document_identifier     = kwargs.get('document_identifier', None)
    document_keywords       = kwargs.get('document_keywords', None)
    document_language       = kwargs.get('document_language', None)
    document_lastModifiedBy = kwargs.get('document_lastModifiedBy', None)
    document_lastPrinted    = kwargs.get('document_lastPrinted', None)
    document_modified       = kwargs.get('document_modified', None)
    document_revision       = kwargs.get('document_revision', None)
    document_subject        = kwargs.get('document_subject', None)
    document_title          = kwargs.get('document_title', None)
    document_version        = kwargs.get('document_version', None)
    
    '''
    print(' ' * 12 + 'document author: ' +  document_author)
    print(' ' * 12 + 'book subject: ' +  book_subject)
    print(' ' * 12 + 'book author: ' +  book_author)
    print(' ' * 12 + 'book company: ' +  book_company)
    print(' ' * 12 + 'book created: ' +  str(book_created))
    print(' ' * 12 + 'book comments: ' +  book_comments)
    print(' ' * 12 + 'sheet title: ' +  sheet_title)
    '''
    
    #открываем шаблон и используем его в качестве рабочего документа
    document = Document(os.path.join(script_dirName, template_address))

    #заполняем свойства документа
    if document_author          != None: document.core_properties.author            = document_author
    if document_category        != None: document.core_properties.category          = document_category
    if document_comments        != None: document.core_properties.comments          = document_comments
    if document_contentStatus   != None: document.core_properties.content_status    = document_contentStatus
    if document_created         != None: document.core_properties.created           = document_created
    if document_identifier      != None: document.core_properties.identifier        = document_identifier
    if document_keywords        != None: document.core_properties.keywords          = document_keywords
    if document_language        != None: document.core_properties.language          = document_language
    if document_lastModifiedBy  != None: document.core_properties.last_modified_by  = document_lastModifiedBy
    if document_lastPrinted     != None: document.core_properties.last_printed      = document_lastPrinted
    if document_modified        != None: document.core_properties.modified          = document_modified
    if document_revision        != None: document.core_properties.revision          = document_revision
    if document_subject         != None: document.core_properties.subject           = document_subject
    if document_title           != None: document.core_properties.title             = document_title
    if document_version         != None: document.core_properties.version           = document_version

    #индексы столбцов в таблице перечня
    pe3Table_colIndex_designator = 0
    pe3Table_colIndex_label      = 1
    pe3Table_colIndex_quantity   = 2
    pe3Table_colIndex_annotation = 3

    #массивы таблиц
    table_titleBlock_main  = [document.tables[0].cell(1,1).tables[0], document.tables[1].cell(1,1).tables[0]] #таблица с главным блоком основной надписи
    table_titleBlock_extra = [document.tables[0].cell(0,0).tables[0], document.tables[1].cell(0,0).tables[0]] #таблица с дополнительным блоком основной надписи
    table_pe3Data          = [document.tables[0].cell(0,1).tables[0], document.tables[1].cell(0,1).tables[0]] #таблица с записями перечня элементов

    #заполняем поля основной надписи
    if data.titleblock is not None:
        print('INFO >> Filling title block', end ="... ", flush = True)
        if isinstance(data.titleblock, dict):
            #--- первый лист
            table_titleBlock_main[0].cell( 7,  5).paragraphs[0].text = data.titleblock.get('01a_product_name', '')
            table_titleBlock_main[0].cell( 7,  5).paragraphs[1].text = data.titleblock.get('01b_document_type', '')
            table_titleBlock_main[0].cell( 2, 10).paragraphs[0].text = data.titleblock.get('02_document_designator', '')
            table_titleBlock_main[0].cell( 4,  6).paragraphs[0].text = data.titleblock.get('04_letter_left', '')
            table_titleBlock_main[0].cell( 4,  7).paragraphs[0].text = data.titleblock.get('04_letter_middle', '')
            table_titleBlock_main[0].cell( 4,  8).paragraphs[0].text = data.titleblock.get('04_letter_right', '')
            #table_titleBlock_main[0].cell( 4,  9).paragraphs[0].text = '1' #порядковый номер листа (на документах, состоящих из одного листа, графу не заполняют)
            table_titleBlock_main[0].cell( 7, 10).paragraphs[0].text = data.titleblock.get('09_organization', '')
            table_titleBlock_main[0].cell( 5,  1).paragraphs[0].text = data.titleblock.get('10d_activityType_extra', '')
            table_titleBlock_main[0].cell( 3,  2).paragraphs[0].text = data.titleblock.get('11a_name_designer', '')
            table_titleBlock_main[0].cell( 4,  2).paragraphs[0].text = data.titleblock.get('11b_name_checker', '')
            table_titleBlock_main[0].cell( 5,  2).paragraphs[0].text = data.titleblock.get('11d_name_extra', '')
            table_titleBlock_main[0].cell( 6,  2).paragraphs[0].text = data.titleblock.get('11e_name_normativeSupervisor', '')
            table_titleBlock_main[0].cell( 7,  2).paragraphs[0].text = data.titleblock.get('11f_name_approver', '')
            table_titleBlock_main[0].cell( 3,  4).paragraphs[0].text = data.titleblock.get('13a_signatureDate_designer', '')
            table_titleBlock_main[0].cell( 4,  4).paragraphs[0].text = data.titleblock.get('13b_signatureDate_checker', '')
            table_titleBlock_main[0].cell( 5,  4).paragraphs[0].text = data.titleblock.get('13d_signatureDate_extra', '')
            table_titleBlock_main[0].cell( 6,  4).paragraphs[0].text = data.titleblock.get('13e_signatureDate_normativeSupervisor', '')
            table_titleBlock_main[0].cell( 7,  4).paragraphs[0].text = data.titleblock.get('13f_signatureDate_approver', '')
            table_titleBlock_extra[0].cell(7,  1).paragraphs[0].text = data.titleblock.get('19_original_inventoryNumber', '')
            table_titleBlock_extra[0].cell(5,  1).paragraphs[0].text = data.titleblock.get('21_replacedOriginal_inventoryNumber', '')
            table_titleBlock_extra[0].cell(4,  1).paragraphs[0].text = data.titleblock.get('22_duplicate_inventoryNumber', '')
            table_titleBlock_extra[0].cell(1,  1).paragraphs[0].text = data.titleblock.get('24_baseDocument_designator', '')
            table_titleBlock_extra[0].cell(0,  1).paragraphs[0].text = data.titleblock.get('25_firstReferenceDocument_designator', '')
            #--- второй лист
            table_titleBlock_main[1].cell( 2,  5).paragraphs[0].text = data.titleblock.get('02_document_designator', '')
            table_titleBlock_extra[1].cell(7,  1).paragraphs[0].text = data.titleblock.get('19_original_inventoryNumber', '')
            table_titleBlock_extra[1].cell(5,  1).paragraphs[0].text = data.titleblock.get('21_replacedOriginal_inventoryNumber', '')
            table_titleBlock_extra[1].cell(4,  1).paragraphs[0].text = data.titleblock.get('22_duplicate_inventoryNumber', '')
        else:
            raise ValueError("Invalid titleblock data.")

    #копируем второй лист в качестве шаблона для последующих
    nextPageTemplate = deepcopy(document.tables[1]._tbl)
    print('done.')

    #получаем данные о форматировании таблицы с перечнем элементов (подразумеваем что все строки таблицы имеют одинаковое форматирование)
    pe3Table_cellWidth_designator = getCellEffectiveWidth(table_pe3Data[0].cell(1, pe3Table_colIndex_designator))
    pe3Table_cellWidth_label      = getCellEffectiveWidth(table_pe3Data[0].cell(1, pe3Table_colIndex_label))
    pe3Table_cellWidth_quantity   = getCellEffectiveWidth(table_pe3Data[0].cell(1, pe3Table_colIndex_quantity))
    pe3Table_cellWidth_annotation = getCellEffectiveWidth(table_pe3Data[0].cell(1, pe3Table_colIndex_annotation))
    pe3Table_font_designator = getFont(table_pe3Data[0].cell(1, pe3Table_colIndex_designator).paragraphs[0])
    pe3Table_font_label      = getFont(table_pe3Data[0].cell(1, pe3Table_colIndex_label).paragraphs[0])
    pe3Table_font_quantity   = getFont(table_pe3Data[0].cell(1, pe3Table_colIndex_quantity).paragraphs[0])
    pe3Table_font_annotation = getFont(table_pe3Data[0].cell(1, pe3Table_colIndex_annotation).paragraphs[0])

    print('INFO >> Filling pe3 data', end = "... ", flush = True)
    #заполняем сам перечень элементов
    pe3Table_pageIndex       = 0                        #индекс текущей страницы
    pe3Table_rowIndex_start  = 1                        #индекс начальной строки в таблице перечня на листе
    pe3Table_rowIndex        = pe3Table_rowIndex_start  #индекс текущей строки в таблице на листе
    pe3Table_rowIndex_end    = 29                       #индекс конечной строки в таблице перечня на листе
    
    print('page: ' + str(pe3Table_pageIndex + 1), end = ", ", flush = True)
    pe3_rowIndex = 0
    while pe3_rowIndex < len(data.rows):
        if data.rows[pe3_rowIndex].flag == data.Row.FlagType.SPACER:
            #если строка-отступ то ничего не пишем, а если первая на странице то и индекс не увеличиваем
            pe3_rowIndex += 1
            if pe3Table_rowIndex != pe3Table_rowIndex_start:
                pe3Table_rowIndex += 1

        elif data.rows[pe3_rowIndex].flag == data.Row.FlagType.TITLE:
            if pe3Table_rowIndex > pe3Table_rowIndex_end - 1:
                #если заголовок дальше определённого уровня до конца страницы то пропускаем запись (ждём до начала новой страницы)
                pass
            else:
                table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_label).paragraphs[0].text = data.rows[pe3_rowIndex].label
                table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_label).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                pe3_rowIndex += 1
            pe3Table_rowIndex += 1
        else:
            #записываем строку перечня элементов
            #--- формируем строки перечня
            pe3Table_text_designator = data.rows[pe3_rowIndex].designator.split('\n')
            pe3Table_text_label      = data.rows[pe3_rowIndex].label.split('\n')
            pe3Table_text_quantity   = [str(data.rows[pe3_rowIndex].quantity)]
            pe3Table_text_annotation = data.rows[pe3_rowIndex].annotation.split('\n')

            if pe3Table_wordwrap_desig:
                for i in range(len(pe3Table_text_designator)):
                    lines = wrapText(pe3Table_text_designator[i], pe3Table_cellWidth_designator, pe3Table_font_designator, pe3Table_wordwrap_delimiter_desig)
                    pe3Table_text_designator[i] = lines[0]
                    for j in range(1, len(lines)):
                        pe3Table_text_designator.insert(i + j, lines[j])
            
            if pe3Table_wordwrap_label:
                for i in range(len(pe3Table_text_label)):
                    lines = wrapText(pe3Table_text_label[i], pe3Table_cellWidth_label, pe3Table_font_label, pe3Table_wordwrap_delimiter_label)
                    pe3Table_text_label[i] = lines[0]
                    for j in range(1, len(lines)):
                        pe3Table_text_label.insert(i + j, lines[j])

            if pe3Table_wordwrap_annot:
                for i in range(len(pe3Table_text_annotation)):
                    lines = wrapText(pe3Table_text_annotation[i], pe3Table_cellWidth_annotation, pe3Table_font_annotation, pe3Table_wordwrap_delimiter_annot)
                    pe3Table_text_annotation[i] = lines[0]
                    for j in range(1, len(lines)):
                        pe3Table_text_annotation.insert(i + j, lines[j])
            
            #--- выравниваем размеры массивов со строками
            pe3Table_entryHeight = max(len(pe3Table_text_designator), len(pe3Table_text_label), len(pe3Table_text_quantity), len(pe3Table_text_annotation)) #количество строк таблицы перечня для данной записи
            pe3Table_text_designator[:] += ['']*(pe3Table_entryHeight - len(pe3Table_text_designator))
            pe3Table_text_label[:] += ['']*(pe3Table_entryHeight - len(pe3Table_text_label))
            pe3Table_text_quantity[:] += ['']*(pe3Table_entryHeight - len(pe3Table_text_quantity))
            pe3Table_text_annotation[:] += ['']*(pe3Table_entryHeight - len(pe3Table_text_annotation))

            #--- проверяем влезает ли запись на текущую страницу
            if pe3Table_rowIndex + pe3Table_entryHeight - 1 <= pe3Table_rowIndex_end:
                #запись влезает - записываем строки в таблицу
                for i in range(pe3Table_entryHeight):
                    table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_designator).paragraphs[0].text = pe3Table_text_designator[i]
                    table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_label).paragraphs[0].text = pe3Table_text_label[i]
                    table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_quantity).paragraphs[0].text = pe3Table_text_quantity[i]
                    table_pe3Data[pe3Table_pageIndex].cell(pe3Table_rowIndex, pe3Table_colIndex_annotation).paragraphs[0].text = pe3Table_text_annotation[i]
                    pe3Table_rowIndex += 1
                pe3_rowIndex += 1
            else:
                #запись не влезает - делаем индекс больше максимального (чтобы добавилась новая страница) и прогоняем запись поновой
                pe3Table_rowIndex = pe3Table_rowIndex_end + 1
            
        #проверяем заполнился ли текущий лист
        if pe3Table_rowIndex > pe3Table_rowIndex_end:
            pe3Table_pageIndex += 1
            print(str(pe3Table_pageIndex + 1), end =", ", flush = True)
            pe3Table_rowIndex = pe3Table_rowIndex_start
            pe3Table_rowIndex_end = 32
            if pe3Table_pageIndex > 1:
                #если третий лист то добавлять параграф не надо так как в шаблоне не удаляется лишний параграф в самом конце
                if pe3Table_pageIndex == 2:
                    nextParagraph = document.paragraphs[1]
                else:
                    nextParagraph = document.add_paragraph()
                #добавляем следующий лист в документ
                nextParagraph._p.addnext(deepcopy(nextPageTemplate))
            
                #добавляем новый лист в массивы таблиц
                table_titleBlock_main.append(document.tables[pe3Table_pageIndex].cell(1,1).tables[0])
                table_titleBlock_extra.append(document.tables[pe3Table_pageIndex].cell(0,0).tables[0])
                table_pe3Data.append(document.tables[pe3Table_pageIndex].cell(0,1).tables[0])
                
            #ставим номер листа
            table_titleBlock_main[pe3Table_pageIndex].cell(2,6).tables[0].cell(1,0).paragraphs[0].text = str(pe3Table_pageIndex + 1)

    #записываем общее количество листов в документе
    table_titleBlock_main[0].cell(4,10).paragraphs[0].text = str(pe3Table_pageIndex + 1)
    print('done.') 

    #смотрим количество листов
    if pe3Table_pageIndex > 0:
        #ставим номер на первый лист если их больше одного
        table_titleBlock_main[0].cell( 4,  9).paragraphs[0].text = '1'
    else:
        #удаляем второй лист
        document.tables[1]._element.getparent().remove(document.tables[1]._element)
        document.paragraphs[1]._element.getparent().remove(document.paragraphs[1]._element)
    
    print('INFO >> Saving result', end ="... ", flush = True) 
    document.save(os.path.join(script_dirName, address))
    print('ok.')
    print('INFO >> pe3-docx export completed.')   

#Возвращает параметры шрифта [имя, размер, жирный, курсив] параграфа
def getFontParams(paragraph):
    #идём вверх по дереву стилей пока не найдём заполненное поле имени шрифта
    style = paragraph.style
    while style.font.name == None:
        if style.base_style == None:
            break
        style = style.base_style
    family = style.font.name

    #идём вверх по дереву стилей пока не найдём заполненное поле размера шрифта
    style = paragraph.style
    while style.font.size == None:
        if style.base_style == None:
            break
        style = style.base_style
    size = int(style.font.size/Pt(1))

    #идём вверх по дереву стилей пока не найдём заполненное поле жирности шрифта или нет родительского стиля
    style = paragraph.style
    while style.font.bold == None:
        if style.base_style == None:
            break
        style = style.base_style
    if style.font.bold == None:
        bold = False
    else:
        bold = style.font.bold

    #идём вверх по дереву стилей пока не найдём заполненное поле курсива шрифта или нет родительского стиля
    style = paragraph.style
    while style.font.italic == None:
        if style.base_style == None:
            break
        style = style.base_style
    if style.font.italic == None:
        italic = False
    else:
        italic = style.font.italic

    return [family, size, bold, italic]

#Возвращает объект шрифта параграфа
def getFont(paragraph):
    #большая проблема соотнести название шрифта указанного в шаблоне с конкретным файлом шрифта в системе (если он вообще установлен!)
    #пока просто закидываем файл шрифта вручную, может быть когда-нибудь напишется и этот кусок
    #https://python-utilities.readthedocs.io/en/latest/sysfont.html
    #https://stackoverflow.com/questions/43601276/get-font-file-from-font-family
    #https://stackoverflow.com/questions/4190667/how-to-get-width-of-a-truetype-font-character-in-1200ths-of-an-inch-with-python 
    fontParams = getFontParams(paragraph)
    font = ImageFont.truetype("opengosta.otf", fontParams[1])
    return font

#Возвращает эффективную ширину ячейки в миллиметрах (ширину доступную для текста т.е заданная ширина за вычетом полей)
def getCellEffectiveWidth(cell):
    #проблемы с получением значения ширины полей для ячейки, поэтому задаём поля вручную
    #также размеры таблицы должны быть жёстко заданы вручную
    cellWidth = cell.width / Mm(1.0)
    margin_start = 1.0
    margin_end   = margin_start
    return cellWidth - margin_start - margin_end

#Разбивает текст на строки (ширина в миллиметрах)
def wrapText(text, width, font, resplitChars = ' '):
    max_width = width / (25.4/72)
    if getTextWidth(font, text) > max_width:
        #разбиваем текст на блоки по заданным разделителям и приклеиваем разделитель обратно к концу блока
        blocks = re.split('([' + resplitChars + '])', text)
        for i in range(len(blocks) // 2):
            blocks[i] = blocks[i] + blocks[i + 1]
            blocks.pop(i + 1)
        lines = [[]]
        for block in blocks:
            #try putting this block in last line then measure
            lines[-1].append(block)
            line = ''.join(lines[-1])
            if getTextWidth(font, line) > max_width: #too wide
                #take it back out, put it on the next line, then measure again
                if len(lines[-1]) > 1:   #take back out only if more than one block
                    lines.append([lines[-1].pop()])
        return [''.join(line) for line in lines]
    else:
        return [text]

#Вычисляет ширину текста
def getTextWidth(font, text):
    bbox = font.getbbox(text)
    return bbox[2] - bbox[0]

#Записывает адреса ячеек в таблицы шаблона
def indexTemplate(resultAddress, templateAddress = None):
    if templateAddress == None:
        templateAddress = template_defaultAddress
    
    document = Document(os.path.join(script_dirName, templateAddress))

    for page_index in range(len(document.tables)):
        nestedTables = (document.tables[page_index].cell(0,0).tables[0], document.tables[page_index].cell(0,1).tables[0], document.tables[page_index].cell(1,1).tables[0])
        for table_index in range(len(nestedTables)):
            for row_index in range(len(nestedTables[table_index].rows)):
                for column_index in range(len(nestedTables[table_index].columns)):
                    nestedTables[table_index].cell(row_index,column_index).text = str(table_index) + ':' + str(row_index) + ',' + str(column_index) 

    document.save(os.path.join(script_dirName, resultAddress))

